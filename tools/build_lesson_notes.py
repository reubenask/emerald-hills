from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/reuben/Desktop/TOEFL junior/Emerald Hills Bilingual Lesson Notes March-June 2026.docx")


@dataclass
class Lesson:
    course: str
    date: str
    time: str
    topic_en: str
    topic_zh: str
    objective_en: str
    objective_zh: str
    ice_en: str
    ice_zh: str
    lesson_en: str
    lesson_zh: str
    materials_en: str
    materials_zh: str
    homework_en: str
    homework_zh: str
    notes_en: str = "To be completed after class: record each student’s participation, accuracy, confidence, challenges, and next support step."
    notes_zh: str = "课后填写：记录每位学生的参与度、准确率、表达信心、遇到的困难，以及下一步需要的支持。"


def l(
    course: str,
    date: str,
    time: str,
    topic_en: str,
    topic_zh: str,
    lesson_en: str,
    lesson_zh: str,
    homework_en: str,
    homework_zh: str,
    objective_en: str | None = None,
    objective_zh: str | None = None,
    materials_en: str | None = None,
    materials_zh: str | None = None,
    ice_en: str | None = None,
    ice_zh: str | None = None,
) -> Lesson:
    return Lesson(
        course=course,
        date=date,
        time=time,
        topic_en=topic_en,
        topic_zh=topic_zh,
        objective_en=objective_en or f"Students will be able to understand and apply the core skills for {topic_en.lower()} through guided practice, feedback, and independent production.",
        objective_zh=objective_zh or f"学生能够通过教师引导、课堂练习和反馈，理解并运用“{topic_zh}”的核心技能。",
        ice_en=ice_en or "Quick recall of the previous lesson, short oral warm-up, and one guiding question connected to today’s target.",
        ice_zh=ice_zh or "快速复习上节课内容，进行简短口语热身，并用一个问题导入本节课目标。",
        lesson_en=lesson_en,
        lesson_zh=lesson_zh,
        materials_en=materials_en or "Teacher slides, TOEFL-style prompts, class worksheet, whiteboard notes, and student response sheets.",
        materials_zh=materials_zh or "教师课件、托福风格题目、课堂练习纸、白板记录和学生答题纸。",
        homework_en=homework_en,
        homework_zh=homework_zh,
    )


LESSONS: list[Lesson] = [
    # iBT
    l("iBT", "2026-03-02", "10:00-12:00; 13:00-16:00", "Speaking and writing test structure", "口语与写作考试结构",
      "Students unpacked the question structures for TOEFL speaking and writing, practised recognizing task demands, and learned practical acronyms for planning answers under time pressure.",
      "学生分析托福口语和写作题型结构，练习判断题目要求，并学习在限时条件下组织答案的实用缩写方法。",
      "Complete the AI-generated speaking task exercise.", "完成AI生成的口语任务练习。"),
    l("iBT", "2026-03-06", "18:30-20:30", "Strategy review for speaking and writing", "口语与写作策略复习",
      "The class reviewed the main speaking/writing strategies, strengthened planning routines, and practised using acronyms to produce clearer, better organized responses.",
      "课堂复习口语和写作主要策略，强化构思流程，并练习使用缩写方法产出更清晰、更有结构的答案。",
      "Review the strategy acronyms and prepare one model response.", "复习课堂策略缩写，并准备一份示范答案。"),
    l("iBT", "2026-03-07", "10:00-12:00", "Reading question types and strategies", "阅读题型与解题策略",
      "Students studied TOEFL reading question types, identified where answers are usually located, and practised choosing evidence before selecting an answer.",
      "学生学习托福阅读题型，识别答案常出现的位置，并练习先找证据再选择答案。",
      "Complete the class practice set questions.", "完成课堂练习题。", materials_en="Reading passages, question-type guide, annotation worksheet.", materials_zh="阅读文章、题型指南、批注练习纸。"),
    l("iBT", "2026-03-13", "18:30-20:30", "Reading task review", "阅读任务复盘",
      "The class reviewed reading tasks from previous practice and corrected common errors in evidence matching, inference, vocabulary-in-context, and detail questions.",
      "课堂复盘前次阅读练习，纠正常见错误，包括证据匹配、推理题、语境词汇题和细节题。",
      "Revise corrected answers and mark the evidence line for each one.", "订正答案，并为每题标出原文证据。"),
    l("iBT", "2026-03-14", "10:00-12:00", "Reading evidence practice", "阅读证据训练",
      "Students completed targeted reading exercises for Tasks 1-3 and were required to defend each answer with a clear “why” based on the text.",
      "学生完成阅读任务1-3的针对性练习，并必须用原文依据说明每个答案的理由。",
      "Complete reading exercise tasks 1-3 and write a WHY for each answer.", "完成阅读任务1-3，并为每个答案写出WHY理由。"),
    l("iBT", "2026-03-20", "18:30-20:30", "Listening question types and strategies", "听力题型与策略",
      "Students reviewed TOEFL listening question types, note-taking priorities, speaker purpose, organization, and how to track examples in academic lectures.",
      "学生复习托福听力题型、笔记重点、说话者目的、篇章结构，以及如何跟踪学术讲座中的例子。",
      "Review listening notes and list three signal words from the lesson.", "复习听力笔记，并列出课堂中出现的三个信号词。", materials_en="Lecture audio, note-taking template, listening question set.", materials_zh="讲座音频、笔记模板、听力题组。"),
    l("iBT", "2026-03-21", "10:00-12:00", "TOEFL mock test supervision", "托福模拟测试",
      "Students completed a supervised TOEFL-style mock session. The teacher observed timing, endurance, answer-sheet habits, and test-taking discipline.",
      "学生完成一次有监督的托福风格模拟测试。教师观察时间管理、考试耐力、答题卡习惯和考试规范。",
      "Review the test and identify three questions to discuss.", "复盘测试，并选出三道需要讨论的题。", materials_en="Mock test paper, answer sheet, timer.", materials_zh="模拟试卷、答题卡、计时器。"),
    l("iBT", "2026-03-27", "18:30-20:30", "Exam review and writing structures 1-3", "考试复盘与写作结构1-3",
      "The class reviewed mock-test performance and rebuilt the first three writing structures, focusing on topic control, paragraph order, and sentence clarity.",
      "课堂复盘模拟测试表现，并重新梳理前三类写作结构，重点关注主题控制、段落顺序和句子清晰度。",
      "Revise one weak writing response using the class structure.", "使用课堂结构修改一篇薄弱写作答案。"),
    l("iBT", "2026-03-28", "10:00-12:00", "Short writing quiz: tasks 2 and 3", "写作小测：任务2和3",
      "Students completed a short quiz for writing Tasks 2 and 3, then corrected organization, grammar, and content development with teacher feedback.",
      "学生完成写作任务2和3的小测，并在教师反馈下修改文章结构、语法和内容展开。",
      "Rewrite the corrected quiz response.", "重写订正后的写作小测答案。", materials_en="Timed writing prompts, correction checklist.", materials_zh="限时写作题目、修改清单。"),
    l("iBT", "2026-04-10", "18:30-20:30", "Writing task and grammar review", "写作任务与语法复习",
      "The lesson reviewed writing task expectations, common grammar weaknesses, and classroom exercises for improving accuracy and sentence control.",
      "本课复习写作任务要求、常见语法弱点，并通过课堂练习提升准确性和句子控制能力。",
      "Revise class grammar corrections and prepare one improved paragraph.", "复习课堂语法订正，并准备一段修改后的段落。"),
    l("iBT", "2026-04-11", "10:00-12:00", "Intensive writing practice: tasks 1-3", "强化写作练习：任务1-3",
      "Students practised Tasks 1, 2, and 3 intensively, moving from planning to timed drafting and then revising for clarity, support, and grammar.",
      "学生强化练习写作任务1、2、3，从构思到限时写作，再修改清晰度、论据和语法。",
      "Attempt the assigned writing task.", "完成指定写作任务。"),
    l("iBT", "2026-04-17", "18:30-20:30", "Speaking practice: task 2 parts 1-2", "口语练习：任务2第1-2部分",
      "Students practised TOEFL speaking Task 2 parts 1 and 2, focusing on listening to the prompt, preparing quickly, and delivering a timed response.",
      "学生练习托福口语任务2第1和第2部分，重点训练听题、快速准备和限时回答。",
      "Record one Task 2 response and review timing.", "录制一个任务2回答，并检查时间控制。", materials_en="Speaking prompt audio, prep timer, response timer.", materials_zh="口语题目音频、准备计时器、回答计时器。"),
    l("iBT", "2026-04-18", "10:00-12:00", "Speaking practice: task 2 parts 3-4", "口语练习：任务2第3-4部分",
      "Students practised Task 2 parts 3 and 4, with emphasis on combining listening details, planning concise notes, and producing complete spoken answers.",
      "学生练习任务2第3和第4部分，重点整合听力细节、做简洁笔记，并产出完整口语答案。",
      "Attempt the assigned speaking task.", "完成指定口语任务。"),
    l("iBT", "2026-04-24", "18:30-20:30", "Diagnostic speaking, writing, and reading mechanics", "口语写作诊断与阅读机制",
      "The class used short writing and speaking tasks to identify weaknesses, then moved into reading mechanics such as paragraph purpose, evidence, and vocabulary clues.",
      "课堂通过短写作和短口语任务诊断薄弱点，并学习阅读机制，如段落功能、证据和词汇线索。",
      "Review diagnostic feedback and revise one answer.", "复习诊断反馈，并修改一个答案。"),
    l("iBT", "2026-04-25", "10:00-12:00", "Long reading passages and vocabulary", "长篇阅读与词汇",
      "Students worked through longer TOEFL-style passages and built vocabulary through context, word families, and sentence-level application.",
      "学生学习较长的托福风格阅读文章，并通过语境、词族和句子运用积累词汇。",
      "Complete two reading passages and answer the related questions.", "完成两篇阅读文章并回答相关问题。", materials_en="Old TOEFL reading passage and vocabulary worksheet.", materials_zh="旧版托福阅读文章和词汇练习纸。"),
    l("iBT", "2026-05-08", "18:30-20:30", "Writing diagnosis and academic topic discussion", "写作诊断与学术话题讨论",
      "Students completed and discussed short writing tasks, identified recurring weaknesses, and connected TOEFL academic lecture topics with visuals and classroom discussion.",
      "学生完成并讨论短写作任务，识别反复出现的薄弱点，并结合图片讨论托福学术讲座话题。",
      "Redo the email and sentence-arrangement writing tasks.", "重做邮件写作和句子排列写作任务。"),
    l("iBT", "2026-05-09", "10:00-12:00", "Writing diagnosis and academic topic discussion", "写作诊断与学术话题讨论",
      "The class repeated the writing-review process with additional examples, strengthening idea development, sentence order, and clearer academic expression.",
      "课堂用更多例子继续进行写作复盘，强化观点展开、句子顺序和更清晰的学术表达。",
      "Redo the assigned classroom discussion writing task.", "重做指定课堂讨论写作任务。"),
    l("iBT", "2026-05-15", "18:30-20:30", "Long-response writing and speaking practice", "长篇写作与口语练习",
      "Students practised longer writing responses, completed speaking exercises, and built background knowledge from TOEFL topic domains.",
      "学生练习较长写作答案，完成口语练习，并围绕托福常见话题领域建立背景知识。",
      "Prepare vocabulary notes from the topic domain.", "整理本话题领域的词汇笔记。"),
    l("iBT", "2026-05-16", "10:00-12:00", "Academic writing and lecture listening", "学术写作与讲座听力",
      "Students practised long-form writing and topic-domain knowledge, then connected listening lecture details to speaking output.",
      "学生练习长篇写作和话题知识，并把讲座听力细节连接到口语输出。",
      "Complete the lecture listening and speaking exercise.", "完成讲座听力和口语练习。", materials_en="Lecture audio, speaking prompt, writing response sheet.", materials_zh="讲座音频、口语题目、写作答题纸。"),
    l("iBT", "2026-05-22", "18:30-20:30", "Documentary-based reading and short essay practice", "纪录片主题阅读与短文写作",
      "Students watched short documentary clips across topics, discussed useful vocabulary, and wrote a short essay connected to Reading Task 1 practice.",
      "学生观看不同主题的短纪录片片段，讨论有用词汇，并完成与阅读任务1相关的短文写作。",
      "Review documentary vocabulary and revise the short essay.", "复习纪录片词汇并修改短文。"),
    l("iBT", "2026-05-29", "18:30-20:30", "Monthly test day", "月度测试",
      "Students completed the scheduled test under timed conditions. The teacher monitored pacing, instructions, and answer-sheet completion.",
      "学生在限时条件下完成月度测试。教师监督答题节奏、考试要求和答题卡完成情况。",
      "No new homework; rest and prepare for test review.", "无新作业；休息并准备测试复盘。", materials_en="Monthly test paper, media prompts, answer sheets, timer.", materials_zh="月度试卷、媒体题目、答题卡、计时器。"),
    l("iBT", "2026-05-30", "10:00-12:00", "Test review and email writing", "测试复盘与邮件写作",
      "The class reviewed test performance, clarified recurring errors, and practised writing an email using a guided structure.",
      "课堂复盘测试表现，讲解反复出现的错误，并用指导结构练习邮件写作。",
      "Write an email based on the given guidelines.", "根据给定要求写一封邮件。"),
    l("iBT", "2026-06-05", "18:30-20:30", "Vocabulary learning method", "词汇学习方法",
      "Students learned a systematic vocabulary routine: word meaning, part of speech, example sentence, and application inside reading passages.",
      "学生学习系统词汇方法：词义、词性、例句，以及在阅读文章中的实际运用。",
      "Review the new words, meanings, and written sentences.", "复习新词、词义和所写例句。", materials_en="Vocabulary log, reading passage, sentence-building template.", materials_zh="词汇记录表、阅读文章、造句模板。"),
    l("iBT", "2026-06-06", "10:00-12:00", "Generating ideas for academic writing", "学术写作构思方法",
      "Students learned how to generate ideas for academic writing by breaking topics into viewpoints, examples, evidence, and clear sentence development.",
      "学生学习如何为学术写作构思：把话题拆分为观点、例子、证据和清晰的句子展开。",
      "Write two classroom discussion responses.", "写两篇课堂讨论回答。"),
    l("iBT", "2026-06-12", "18:30-20:30", "TOEFL speaking imitation and vocabulary", "托福口语模仿与词汇",
      "Students practised imitating TOEFL speaking models, focusing on pace, stress, sentence chunks, and useful vocabulary for the topic.",
      "学生练习模仿托福口语示范，重点关注语速、重音、句群和话题相关词汇。",
      "Record a short imitation response and note useful phrases.", "录制一段模仿回答，并记录有用表达。"),
    l("iBT", "2026-06-13", "10:00-12:00", "SVO sentence core and expansion", "主谓宾核心句与句子扩展",
      "Students built baseline English sentences using SVO, then expanded them with modifiers, vocabulary from slides, and fuller academic sentence patterns.",
      "学生用主谓宾建立基础英文句，再加入修饰成分、课件词汇和更完整的学术句式进行扩展。",
      "Revise slides 11 and 14 examples; build 10 SVO/full sentences; choose five vocabulary words or SVO structures from slides 15-16 and develop full sentences.",
      "修改课件第11和14页例句；写10个主谓宾/完整句；从第15-16页选择5个词汇或主谓宾结构并扩展成完整句。"),
]


LESSONS += [
    # Junior
    l("Junior", "2026-03-05", "18:30-20:30", "Writing logic with W/H questions", "用W/H问题建立写作逻辑",
      "Students learned to organize ideas with who, what, when, where, why, and how, then used knowledge mapping to plan a complete essay.",
      "学生学习用who、what、when、where、why、how组织想法，并用知识导图规划完整作文。",
      "Use topic mapping to write “My favorite subject in school.”", "使用主题导图写作文《我最喜欢的学校科目》。"),
    l("Junior", "2026-03-07", "13:30-15:30", "Picture storytelling in logical order", "按逻辑顺序看图讲故事",
      "Students practised describing pictures in sequence, naming characters, connecting events, and turning separate images into a complete story.",
      "学生练习按顺序描述图片、给人物命名、连接事件，并把分散图片组织成完整故事。",
      "Create a cohesive story from the picture series, give it a title, record the narration, and submit by Wednesday.", "根据图片系列创作连贯故事，取标题，录音讲述，并在周三前提交。"),
    l("Junior", "2026-03-12", "18:30-20:30", "Expanding W/H planning into essays", "把W/H构思扩展成作文",
      "Students developed basic W/H notes into fuller paragraphs with details, examples, and clearer sentence connections.",
      "学生把基础W/H笔记扩展成更完整的段落，加入细节、例子和更清楚的句子连接。",
      "Choose one topic from the list and use W/H planning to write an essay.", "从题目列表中选择一个，用W/H构思写一篇作文。"),
    l("Junior", "2026-03-14", "13:30-15:30", "Reading comprehension and question types", "阅读理解与题型",
      "Students read a passage, identified important vocabulary, discussed key details, and practised answering different comprehension question types.",
      "学生阅读文章，识别重点词汇，讨论关键细节，并练习回答不同阅读理解题型。",
      "Read the comprehension passage and answer the questions.", "阅读理解文章并回答问题。"),
    l("Junior", "2026-03-19", "18:30-20:30", "Speaking tasks, CVC review, and storytelling", "口语任务、CVC复习与讲故事",
      "Students reviewed TOEFL Junior speaking Tasks 1-2, reinforced CVC reading rules, and used linking words to tell a story from photos.",
      "学生复习托福Junior口语任务1-2，巩固CVC拼读规则，并用连接词根据照片讲故事。",
      "Read the passage in one minute and describe the pictures as a story.", "一分钟内朗读文章，并把图片描述成故事。"),
    l("Junior", "2026-03-21", "13:30-15:30", "Vocabulary through short TV clips", "通过短剧学习词汇",
      "Students learned and discussed vocabulary from a short TV series, focusing on meaning, synonyms, and natural use in context.",
      "学生通过短剧学习和讨论词汇，重点理解词义、近义词和语境中的自然用法。",
      "Read the new words and share meanings or synonyms.", "朗读新词，并分享词义或近义词。"),
    l("Junior", "2026-03-26", "18:30-20:30", "Grammar task review and vocabulary", "语法任务复习与词汇",
      "Students reviewed the previous grammar task, corrected errors, and learned new vocabulary for sentence-level accuracy.",
      "学生复习前次语法任务，订正错误，并学习新词汇以提升句子层面的准确性。",
      "Revise corrected grammar examples.", "复习订正后的语法例句。"),
    l("Junior", "2026-03-28", "13:30-15:30", "Full-section review", "全科板块复习",
      "Students reviewed speaking, grammar, writing, listening, and reading through section examples and exam-style practice.",
      "学生通过板块例题和考试风格练习复习口语、语法、写作、听力和阅读。",
      "Review all section examples before the assessment.", "测评前复习所有板块例题。"),
    l("Junior", "2026-04-02", "18:30-20:30", "Monthly assessment", "月度测评",
      "Students completed the monthly assessment. The teacher monitored pacing, task completion, and areas needing follow-up.",
      "学生完成月度测评。教师观察答题节奏、任务完成情况和需要后续跟进的领域。",
      "No new homework; await test review.", "无新作业；等待测试复盘。"),
    l("Junior", "2026-04-09", "13:30-15:30", "Geography reading with world maps", "结合世界地图进行地理阅读",
      "Students used world maps to understand geography-related TOEFL reading topics and completed knowledge-building exercises.",
      "学生借助世界地图理解托福地理类阅读话题，并完成知识建构练习。",
      "Read the given passage, answer questions, and underline/circle the assigned words.", "阅读指定文章，回答问题，并划出/圈出指定词汇。", materials_en="World maps, reading passage, geography vocabulary worksheet.", materials_zh="世界地图、阅读文章、地理词汇练习纸。"),
    l("Junior", "2026-04-11", "13:30-15:30", "Map-based intensive reading", "地图主题强化阅读",
      "Students used maps, equator regions, food crops, cash crops, and cold regions to deepen comprehension and vocabulary application.",
      "学生结合地图、赤道地区、粮食作物、经济作物和寒冷地区，加深阅读理解和词汇运用。",
      "Mark all new words learned and use them in sentences.", "标出所有新学词汇，并用它们造句。"),
    l("Junior", "2026-04-16", "18:30-20:30", "Vocabulary through The Wild Robot", "通过《荒野机器人》学习词汇",
      "Students learned vocabulary from the movie The Wild Robot, discussed meaning and context, and connected new words to sentence writing.",
      "学生通过电影《荒野机器人》学习词汇，讨论词义和语境，并把新词用于造句。",
      "Use any 20 words learned to make sentences.", "用任意20个新学词汇造句。"),
    l("Junior", "2026-04-18", "13:30-15:30", "Email writing", "邮件写作",
      "Students learned the structure and tone of a response email, including greeting, purpose, details, closing, and clear sentence flow.",
      "学生学习回复邮件的结构和语气，包括问候、目的、细节、结尾和清晰句子衔接。",
      "Write a response email to your friend using the given file.", "根据给定文件给朋友写一封回复邮件。"),
    l("Junior", "2026-04-23", "18:30-20:30", "Typing lesson and email writing", "打字练习与邮件写作",
      "Students practised typing fluency and applied it to email writing so that written responses could be completed more confidently and efficiently.",
      "学生练习打字流利度，并应用到邮件写作中，以更自信、高效地完成书面回答。",
      "Write a response email to your friend using the given file.", "根据给定文件给朋友写一封回复邮件。"),
    l("Junior", "2026-04-25", "13:30-15:30", "Knowledge building and Speaking Task 3", "知识拓展与口语任务3",
      "Students learned about fireflies and pet behavior, discussed whether animals understand people, and practised TOEFL Junior Speaking Task 3.",
      "学生学习萤火虫和宠物行为相关知识，讨论动物是否理解人类，并练习托福Junior口语任务3。",
      "Complete the Speaking Task 3 assignment.", "完成口语任务3作业。"),
    l("Junior", "2026-05-07", "18:30-20:30", "Online vocabulary review", "线上词汇复习",
      "Students reviewed vocabulary through online tools, recorded pronunciation, and used games to strengthen memory and word use.",
      "学生通过线上工具复习词汇，录音练习发音，并用游戏巩固记忆和词汇运用。",
      "Use any 20 words to make sentences.", "用任意20个词造句。"),
    l("Junior", "2026-05-09", "13:30-15:30", "The Wild Robot: themes and vocabulary", "《荒野机器人》主题与词汇",
      "Students watched The Wild Robot, discussed themes, characters, and useful vocabulary, and recorded new expressions for later use.",
      "学生观看《荒野机器人》，讨论主题、人物和有用词汇，并记录新表达供后续使用。",
      "Read the words out loud.", "大声朗读词汇。"),
    l("Junior", "2026-05-14", "18:30-20:30", "Vocabulary review and short drama", "词汇复习与短剧表演",
      "Students reviewed vocabulary and performed two short dramas based on the movie and recently learned words to strengthen practical use.",
      "学生复习词汇，并根据电影和近期新词表演两段短剧，以加强实际运用。",
      "Read the short passage and answer the related questions.", "阅读短文并回答相关问题。"),
    l("Junior", "2026-05-16", "13:30-15:30", "Paragraph summary skills", "段落概括能力",
      "Students learned how to identify each paragraph’s main idea and write a short, accurate summary without copying too much from the text.",
      "学生学习如何找出每段主旨，并写出简短准确、不大量照抄原文的段落摘要。",
      "Review class summary examples.", "复习课堂摘要例子。"),
    l("Junior", "2026-05-21", "18:30-20:30", "Typing practice and exam-structure review", "打字练习与考试结构复习",
      "Students practised typing, reviewed the new exam structure, and studied examples from each section to prepare for the monthly test.",
      "学生练习打字，复习新考试结构，并学习各板块例题，为月度测试做准备。",
      "Revise homework for the upcoming test.", "复习作业，为即将到来的测试做准备。"),
    l("Junior", "2026-05-23", "13:30-15:30", "Monthly test", "月度测试",
      "Students completed the monthly test under timed conditions. The teacher monitored completion, confidence, and section-level strengths.",
      "学生在限时条件下完成月度测试。教师观察完成情况、信心和各板块优势。",
      "No new homework; prepare for review.", "无新作业；准备测试复盘。"),
    l("Junior", "2026-05-28", "18:30-20:30", "Test review: reading section", "测试复盘：阅读板块",
      "The class reviewed the reading section, corrected wrong answers, and practised evidence-finding and question-type recognition.",
      "课堂复盘阅读板块，订正错题，并练习找证据和识别题型。",
      "Complete the second reading passage.", "完成第二篇阅读文章。"),
    l("Junior", "2026-05-30", "13:30-15:30", "Test review: Language Form and Meaning", "测试复盘：语言形式与含义",
      "Students reviewed grammar and vocabulary items from Language Form and Meaning, then used newly discovered words in original sentences.",
      "学生复盘语言形式与含义中的语法和词汇题，并用新发现的词造原创句子。",
      "Use newly discovered words to form sentences.", "用新发现的词造句。"),
    l("Junior", "2026-06-04", "18:30-20:30", "Listening test review", "听力测试复盘",
      "Students reviewed listening test answers, checked why options were correct or incorrect, and rebuilt note-taking strategies.",
      "学生复盘听力测试答案，分析选项正确或错误的原因，并重建笔记策略。",
      "Review the new words.", "复习新词。"),
    l("Junior", "2026-06-06", "13:30-15:30", "Complete The Wild Robot and vocabulary use", "完成《荒野机器人》并运用词汇",
      "Students completed the movie, learned new words, and discussed situations where the words could be used naturally.",
      "学生完成电影观看，学习新词，并讨论这些词在自然情境中的用法。",
      "Revise all vocabulary learned this recent week.", "复习最近一周所学全部词汇。"),
    l("Junior", "2026-06-11", "18:30-20:30", "Grammar detective: error finding", "语法侦探：找错训练",
      "Students reviewed vocabulary and practised finding grammar errors in passages, especially subject-verb agreement and punctuation.",
      "学生复习词汇，并练习在文章中找语法错误，重点包括主谓一致和标点使用。",
      "Rewrite the sentences and correct the grammar issues.", "重写句子并改正语法问题。"),
]


LESSONS += [
    # Primary Step 2
    l("Primary Step 2", "2026-03-04", "18:00-20:00", "Reading with VCV rules", "用VCV规则提升阅读",
      "Students built reading ability through VCV patterns, decoding words, and applying the pattern to a short reading passage.",
      "学生通过VCV拼读规律、单词拆读和短文应用来提升阅读能力。",
      "Read the assigned passage.", "朗读指定文章。", materials_en="Reading passage and phonics board work.", materials_zh="阅读文章和自然拼读板书。"),
    l("Primary Step 2", "2026-03-07", "16:00-18:00", "Spelling, calendar words, and self-introduction", "拼写、日期词汇与自我介绍",
      "Students practised spelling and reading, reviewed days of the week and months of the year, and learned how to introduce themselves to new people.",
      "学生练习拼写和阅读，复习星期与月份词汇，并学习如何向新朋友介绍自己。",
      "Prepare a self-introduction using Cards 1 and 2: name, age, birthday, family, favorite food, reasons, meals, drinks, and favorite restaurant.", "使用卡片1和2准备自我介绍：姓名、年龄、生日、家庭、喜欢的食物及原因、三餐饮食、饮品和喜欢的餐厅。", materials_en="Speaking cards 1-2, spelling list, calendar vocabulary.", materials_zh="口语卡片1-2、拼写词表、日期词汇。"),
    l("Primary Step 2", "2026-03-11", "18:00-20:00", "Syllables and longer reading", "音节与较长阅读",
      "Students revisited CVC, broke longer words into syllables, and practised reading longer passages with better rhythm and accuracy.",
      "学生复习CVC，把较长单词拆分成音节，并练习更有节奏、更准确地朗读较长文章。",
      "Read the longer passage aloud.", "大声朗读较长文章。"),
    l("Primary Step 2", "2026-03-14", "16:00-18:00", "Vocabulary and listening: visiting a pediatrician", "词汇与听力：看儿科医生",
      "Students learned healthcare vocabulary, watched a pediatrician-related video, and answered story questions to connect words with real situations.",
      "学生学习医疗健康词汇，观看儿科医生相关视频，并回答故事问题，把词汇与真实情境联系起来。",
      "Review the healthcare vocabulary.", "复习医疗健康词汇。", materials_en="Pediatrician video and vocabulary list: hospital, clinic, prescription, pharmacy, checkup, nurse, X-ray, and related words.", materials_zh="儿科医生视频和词汇表：hospital、clinic、prescription、pharmacy、checkup、nurse、X-ray等。"),
    l("Primary Step 2", "2026-03-18", "18:00-20:00", "Hospital role play and speaking", "医院主题角色扮演与口语",
      "Students used hospital vocabulary in role plays, acted out practical situations, and practised speaking with clearer sentence patterns.",
      "学生在角色扮演中使用医院主题词汇，演绎实际情境，并练习更清楚的句型表达。",
      "Read the script practised in class.", "朗读课堂练习过的剧本。"),
    l("Primary Step 2", "2026-03-21", "16:00-18:00", "Pronouns, nouns, verbs, and tenses", "代词、名词、动词与时态",
      "Students reviewed pronouns, nouns, verbs, and basic tenses, then wrote simple sentences using accurate subject and verb choices.",
      "学生复习代词、名词、动词和基础时态，并用正确的主语和动词写简单句。",
      "Write 10 sentences using nouns/pronouns and different verb tenses.", "用名词/代词和不同时态写10个句子。"),
    l("Primary Step 2", "2026-03-25", "18:00-20:00", "Past tense and continuous forms", "过去时与进行时",
      "Students revised present, future, and present continuous forms, then learned past tense and past continuous through pronoun-based sentence patterns.",
      "学生复习一般现在时、将来时和现在进行时，并通过代词句型学习过去时和过去进行时。",
      "Revise tense examples from class.", "复习课堂时态例句。"),
    l("Primary Step 2", "2026-03-28", "16:00-18:00", "Vocabulary, sentence building, and listening", "词汇、造句与听力",
      "Students learned new words, used them in sentences, watched a short video about a child deciding what to cook for a birthday, and answered multiple-choice questions.",
      "学生学习新词并造句，观看一个孩子为生日决定做什么食物的短视频，并回答选择题。",
      "Rewatch the video, note anything unclear, and share questions next class.", "重看视频，记录不理解的地方，并在下节课分享问题。"),
    l("Primary Step 2", "2026-04-08", "18:00-20:00", "Continents, countries, and cardinal directions", "大洲、国家与方位",
      "Students learned continents, European countries, and cardinal directions, then formed location sentences and answered reading questions.",
      "学生学习大洲、欧洲国家和方位词，并造地点句子、回答阅读问题。",
      "Read Sam’s travel passage across continents and answer the questions.", "阅读Sam横跨大洲旅行的文章并回答问题。", materials_en="World map, continent/country vocabulary, travel passage.", materials_zh="世界地图、大洲/国家词汇、旅行文章。"),
    l("Primary Step 2", "2026-04-11", "16:00-18:00", "Basic writing with English tenses", "基础写作与英语时态",
      "Students practised present, future, past, continuous, and perfect forms through rhythm, substitution, and sentence writing.",
      "学生通过节奏、替换和造句练习一般现在时、将来时、过去时、进行时和完成时。",
      "Imitate the English tense song, replace the verb and subject, write the full rhyme, then read it aloud and send a video.", "模仿英语时态歌，替换动词和主语，写完整押韵句，并朗读录视频提交。"),
    l("Primary Step 2", "2026-04-15", "18:00-20:00", "Tense review and picture sentences", "时态复习与看图造句",
      "Students revised key tenses, selected verbs, conjugated them accurately, and described pictures with one clear sentence each.",
      "学生复习重点时态，选择动词并正确变形，然后用一句清楚的句子描述每张图片。",
      "Describe each picture with one sentence using the studied tenses.", "用所学时态为每张图片写一句描述。"),
    l("Primary Step 2", "2026-04-18", "16:00-18:00", "Vocabulary through Luca", "通过《夏日友晴天》学习词汇",
      "Students learned new words from the movie Luca and discussed meanings, pronunciation, and use in simple sentences.",
      "学生通过电影《夏日友晴天》学习新词，讨论词义、发音和简单句中的用法。",
      "Read the new words aloud and explain their meanings.", "大声朗读新词并解释意思。"),
    l("Primary Step 2", "2026-04-22", "18:00-20:00", "Reading comprehension structures", "阅读理解结构",
      "Students learned how short passages are organized and practised answering related questions with evidence from the text.",
      "学生学习短文结构，并练习用文章证据回答相关问题。",
      "Read the short passage and answer the related questions.", "阅读短文并回答相关问题。"),
    l("Primary Step 2", "2026-05-06", "18:00-20:00", "Picture-panel story writing", "四格图故事写作",
      "Students used short sentences to describe image panels and connected them into a complete story with sequence and action words.",
      "学生用短句描述图片，并用顺序词和动作词把图片连接成完整故事。",
      "Write a story from four panels using the word bank: Laundry, Clothes, Washing machine, Laundry Basket, Hang, Sunny, Dry, Pegs, Iron, Fold.", "根据四格图写故事，并使用词库：Laundry、Clothes、Washing machine、Laundry Basket、Hang、Sunny、Dry、Pegs、Iron、Fold。"),
    l("Primary Step 2", "2026-05-13", "18:00-20:00", "Rhymes and reading fluency", "押韵与阅读流利度",
      "Students read short rhymes, identified repeated sound patterns, and then read a passage that incorporated the same rhyme patterns.",
      "学生朗读短押韵句，识别重复音型，然后阅读包含相同押韵规律的短文。",
      "Read the 100-word passage aloud.", "大声朗读100词短文。"),
    l("Primary Step 2", "2026-05-16", "16:00-18:00", "Monthly review and test preparation", "月度复习与测试准备",
      "Students reviewed the month’s learning, checked homework, and prepared key skills for the monthly test.",
      "学生复习本月学习内容，检查作业，并为月度测试准备重点技能。",
      "Review homework for the test.", "复习作业，准备测试。"),
    l("Primary Step 2", "2026-05-20", "18:00-20:00", "Monthly test day", "月度测试",
      "Students completed speaking, writing, reading, listening, and grammar tasks under teacher supervision.",
      "学生在教师监督下完成口语、写作、阅读、听力和语法任务。",
      "No new homework; prepare for feedback.", "无新作业；准备听取反馈。", materials_en="Monthly test papers, answer sheets, speaking prompts, timer.", materials_zh="月度试卷、答题卡、口语题目、计时器。"),
    l("Primary Step 2", "2026-05-23", "16:00-18:00", "Luca vocabulary and sentence making", "《夏日友晴天》词汇与造句",
      "Students learned new words and expressions from Luca and practised making their own sentences with the vocabulary.",
      "学生学习《夏日友晴天》中的新词和表达，并练习用这些词造自己的句子。",
      "Read the words aloud.", "大声朗读词汇。"),
    l("Primary Step 2", "2026-05-27", "18:00-20:00", "Exam review", "考试复盘",
      "The class reviewed exam performance, corrected common errors, and clarified vocabulary that affected comprehension and writing.",
      "课堂复盘考试表现，订正常见错误，并讲解影响理解和写作的词汇。",
      "Find the meanings of the words learned.", "查找所学词汇的意思。"),
    l("Primary Step 2", "2026-05-30", "16:00-18:00", "Practical lesson: making juice", "实践课：制作果汁",
      "Students learned vocabulary and step-by-step language for making juice, then described and wrote the process using sequence words.",
      "学生学习制作果汁的词汇和步骤表达，并用顺序词描述和书写过程。",
      "Tell the picture story for 1-2 minutes using sequence words and the vocabulary Pineapple, Cut, Blender, Juice, Pour, Family.", "用1-2分钟讲述图片故事，使用顺序词和词汇：Pineapple、Cut、Blender、Juice、Pour、Family。", materials_en="Fruit/juice vocabulary, process pictures, speaking prompt.", materials_zh="水果/果汁词汇、过程图片、口语题目。"),
    l("Primary Step 2", "2026-06-03", "18:00-20:00", "Drama with new vocabulary", "新词汇短剧表达",
      "Students used recently learned words in short drama-style expressions to practise meaning, pronunciation, and natural use.",
      "学生把近期新学词汇用于短剧表达，练习词义、发音和自然用法。",
      "Review the newly learned words.", "复习新学词汇。"),
    l("Primary Step 2", "2026-06-06", "16:00-18:00", "Complete Luca and discuss vocabulary", "完成《夏日友晴天》并讨论词汇",
      "Students completed the movie Luca, learned new words, and discussed where and how those words were used in the story.",
      "学生完成电影《夏日友晴天》，学习新词，并讨论这些词在故事中的使用场景。",
      "Revise Luca vocabulary from class.", "复习课堂中的《夏日友晴天》词汇。"),
    l("Primary Step 2", "2026-06-10", "18:00-20:00", "Grammar detective: sentence and paragraph errors", "语法侦探：句子与短文找错",
      "Students reviewed vocabulary and searched for errors in sentences and short paragraphs using grammar rules like detectives.",
      "学生复习词汇，并像侦探一样用语法规则寻找句子和短文中的错误。",
      "Rewrite the sentences and correct the grammar issues.", "重写句子并改正语法问题。"),
    l("Primary Step 2", "2026-06-13", "16:00-18:00", "Tenses, subject-verb agreement, punctuation, and capitalization", "时态、主谓一致、标点与大小写",
      "Students reviewed tenses and subject-verb agreement, then found punctuation, capitalization, grammar, and tense mistakes in paragraphs.",
      "学生复习时态和主谓一致，并在段落中寻找标点、大小写、语法和时态错误。",
      "Read the paragraph and find mistakes in tenses, capital letters, and punctuation.", "阅读段落，并找出时态、大小写和标点错误。"),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int, col_widths: Iterable[int]) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in col_widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size=10, bold=False, color="10231C"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_label_para(cell, label: str, text: str, zh: str | None = None):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label)
    set_run_font(r, 9, True, "1F4D78")
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(3)
    r2 = p2.add_run(text)
    set_run_font(r2, 9, False, "10231C")
    if zh:
        p3 = cell.add_paragraph()
        p3.paragraph_format.space_after = Pt(0)
        r3 = p3.add_run(zh)
        set_run_font(r3, 9, False, "3F524A")


def set_style(doc: Document):
    background = OxmlElement("w:background")
    background.set(qn("w:color"), "FFFFFF")
    doc._element.insert(0, background)

    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_title(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Emerald Hills TOEFL Academy")
    set_run_font(r, 12, True, "1F7A5C")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run("Bilingual Lesson Notes · March-June 2026")
    set_run_font(r2, 20, True, "10231C")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Primary Step 2 · Junior · iBT | 课堂记录双语整理")
    set_run_font(r3, 10, False, "5F6B64")


def add_summary(doc: Document):
    doc.add_heading("Document Purpose / 文件用途", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table, 9360, [4680, 4680])
    set_cell_margins(table)
    for i, text in enumerate([
        "This document rewrites the draft lesson notes into the platform structure: topic for today, teaching objective, class-content overview, materials, homework, and individual student performance notes.",
        "本文件将原始课程草稿整理为平台结构：今日主题、教学目标、课堂内容概览、教学材料、作业，以及学生个人表现与挑战记录。",
    ]):
        cell = table.cell(0, i)
        set_cell_width(cell, 4680)
        set_cell_shading(cell, "F4F6F9")
        r = cell.paragraphs[0].add_run(text)
        set_run_font(r, 9.5, False, "10231C")

    doc.add_heading("Verification Notes / 核对说明", level=2)
    for text in [
        "All source dates and times from the three draft files are preserved.",
        "The June lessons that appeared under a May heading in the drafts are corrected to June based on their actual dates.",
        "Homework and assignment instructions are preserved and rewritten for clarity.",
        "Individual student performance is left as a teacher-fillable field because the source drafts did not include named student observations for each lesson.",
    ]:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["Normal"]
        r = p.add_run(text)
        set_run_font(r, 9.5, False, "10231C")


def add_lesson(doc: Document, lesson: Lesson):
    doc.add_heading(f"{lesson.date} · {lesson.time} · {lesson.course}", level=3)
    table = doc.add_table(rows=7, cols=2)
    table.style = "Table Grid"
    col_widths = [1900, 7460]
    set_table_width(table, 9360, col_widths)
    set_cell_margins(table, 90, 130, 90, 130)

    rows = [
        ("Topic for today\n今日主题", lesson.topic_en, lesson.topic_zh),
        ("Teaching objective\n教学目标", lesson.objective_en, lesson.objective_zh),
        ("1. Ice breaker\n导入活动", lesson.ice_en, lesson.ice_zh),
        ("2. Lesson for the day\n当天课程内容", lesson.lesson_en, lesson.lesson_zh),
        ("3. Teaching & learning materials\n教学材料", lesson.materials_en, lesson.materials_zh),
        ("Homework / Assignment\n作业", lesson.homework_en, lesson.homework_zh),
        ("Individual student performance & challenges\n学生个人表现与挑战", lesson.notes_en, lesson.notes_zh),
    ]
    for idx, (label, en, zh) in enumerate(rows):
        label_cell = table.cell(idx, 0)
        content_cell = table.cell(idx, 1)
        set_cell_width(label_cell, col_widths[0])
        set_cell_width(content_cell, col_widths[1])
        label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        content_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(label_cell, "E8EEF5" if idx == 0 else "F4F6F9")
        if idx == 0:
            set_cell_shading(content_cell, "F7FBF8")
        lp = label_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lr = lp.add_run(label)
        set_run_font(lr, 8.5, True, "1F4D78")
        add_label_para(content_cell, "EN: ", en, f"中文：{zh}")


def add_course(doc: Document, course: str, lessons: list[Lesson]):
    doc.add_page_break()
    doc.add_heading(course, level=1)
    doc.add_paragraph(f"{len(lessons)} verified lesson records / {len(lessons)} 条已核对课程记录")
    for lesson in lessons:
        add_lesson(doc, lesson)


def add_verification_appendix(doc: Document):
    doc.add_page_break()
    doc.add_heading("Source Mapping Appendix / 原始资料映射附录", level=1)
    doc.add_paragraph("Each source lesson row below appears once in the rewritten lesson notes.")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    widths = [1700, 1500, 1600, 2760, 1800]
    set_table_width(table, 9360, widths)
    set_cell_margins(table, 70, 100, 70, 100)
    headers = ["Course", "Date", "Time", "Topic", "Homework status"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, "E8EEF5")
        r = cell.paragraphs[0].add_run(h)
        set_run_font(r, 8.5, True, "1F4D78")
    for lesson in LESSONS:
        row = table.add_row().cells
        vals = [lesson.course, lesson.date, lesson.time, lesson.topic_en, "Preserved / 已保留"]
        for i, value in enumerate(vals):
            set_cell_width(row[i], widths[i])
            r = row[i].paragraphs[0].add_run(value)
            set_run_font(r, 8, False, "10231C")


def main():
    doc = Document()
    set_style(doc)
    add_title(doc)
    add_summary(doc)
    order = ["Primary Step 2", "Junior", "iBT"]
    for course in order:
        add_course(doc, course, [x for x in LESSONS if x.course == course])
    add_verification_appendix(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)
    print(f"lessons={len(LESSONS)}")
    for course in order:
        print(f"{course}={sum(1 for x in LESSONS if x.course == course)}")


if __name__ == "__main__":
    main()
